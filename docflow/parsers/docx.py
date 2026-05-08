"""DOCX document parser using python-docx."""

from docflow.parsers.base import BaseParser, ParsedDocument, Section


class DocxParser(BaseParser):
    """Parser for DOCX files using python-docx.

    Extracts text while preserving paragraph styles and table structures,
    creating sections based on heading styles.
    """

    async def parse(self, file_path: str) -> ParsedDocument:
        """Parse a DOCX file into structured content.

        Args:
            file_path: Path to the .docx file.

        Returns:
            ParsedDocument with style-based sections and table content.
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required: pip install python-docx")

        doc = Document(file_path)
        paragraphs = self._extract_paragraphs(doc)
        tables = self._extract_tables(doc)

        full_text_parts: list[str] = []
        sections: list[Section] = []
        current_title = ""
        current_level = 1
        current_lines: list[str] = []
        pos = 0

        for para in paragraphs:
            text = para["text"].strip()
            if not text:
                continue

            if para["style"].startswith("Heading"):
                if current_lines:
                    sections.append(
                        Section(
                            title=current_title,
                            level=current_level,
                            content="\n".join(current_lines),
                            start_char=pos,
                            end_char=pos + sum(len(l) for l in current_lines),
                        )
                    )
                current_title = text
                try:
                    current_level = int(para["style"].replace("Heading ", "").replace("Heading", "1"))
                except ValueError:
                    current_level = 1
                current_lines = []
            else:
                current_lines.append(text)
                pos += len(text) + 1

            full_text_parts.append(text)

        if current_lines:
            sections.append(
                Section(
                    title=current_title,
                    level=current_level,
                    content="\n".join(current_lines),
                    start_char=pos,
                    end_char=pos + sum(len(l) for l in current_lines),
                )
            )

        if tables:
            tables_text = "\n\n".join(tables)
            full_text_parts.append(tables_text)

        full_content = "\n".join(full_text_parts)

        return ParsedDocument(
            content=full_content,
            metadata={"paragraph_count": len(paragraphs), "table_count": len(tables)},
            sections=sections,
            raw_text=full_content,
            title=sections[0].title if sections else "",
            file_type="docx",
        )

    def _extract_paragraphs(self, doc: Any) -> list[dict[str, str]]:
        """Extract paragraphs with their style information.

        Args:
            doc: python-docx Document object.

        Returns:
            List of dicts with text and style for each paragraph.
        """
        paragraphs: list[dict[str, str]] = []
        for para in doc.paragraphs:
            paragraphs.append({"text": para.text, "style": para.style.name})
        return paragraphs

    def _extract_tables(self, doc: Any) -> list[str]:
        """Extract tables as formatted text.

        Args:
            doc: python-docx Document object.

        Returns:
            List of string representations of tables.
        """
        tables: list[str] = []
        for table in doc.tables:
            rows: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            tables.append("\n".join(rows))
        return tables


from typing import Any
