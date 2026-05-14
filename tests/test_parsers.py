"""Tests for document parsers."""

import pytest

from docflow.parsers.markdown import MarkdownParser
from docflow.parsers.html import HTMLParser
from docflow.parsers.csv import CSVParser


class TestMarkdownParser:
    """Tests for the Markdown document parser."""

    @pytest.mark.asyncio
    async def test_parse_markdown(self, sample_dir: "Path") -> None:
        """Test basic markdown parsing without frontmatter."""
        parser = MarkdownParser()
        md_file = sample_dir / "product_docs.md"
        if md_file.exists():
            result = await parser.parse(str(md_file))
            assert result.file_type == "md"
            assert len(result.content) > 0

    @pytest.mark.asyncio
    async def test_parse_markdown_with_frontmatter(self, sample_markdown_content: str, tmp_path: "Path") -> None:
        """Test markdown parsing with YAML frontmatter extraction."""
        md_file = tmp_path / "test.md"
        md_file.write_text(sample_markdown_content)

        parser = MarkdownParser()
        result = await parser.parse(str(md_file))

        assert result.metadata.get("title") == "Test Document"
        assert result.metadata.get("author") == "Test Author"
        assert len(result.content) > 0
        assert result.file_type == "md"

    @pytest.mark.asyncio
    async def test_markdown_sections(self, sample_markdown_content: str, tmp_path: "Path") -> None:
        """Test that markdown sections are correctly extracted from headers."""
        md_file = tmp_path / "test.md"
        md_file.write_text(sample_markdown_content)

        parser = MarkdownParser()
        result = await parser.parse(str(md_file))

        assert len(result.sections) > 0
        section_titles = [s.title for s in result.sections]
        assert any("Introduction" in t for t in section_titles)

    @pytest.mark.asyncio
    async def test_markdown_no_frontmatter(self, tmp_path: "Path") -> None:
        """Test markdown parsing when no frontmatter is present."""
        md_file = tmp_path / "simple.md"
        md_file.write_text("# Simple Document\n\nJust some text.")

        parser = MarkdownParser()
        result = await parser.parse(str(md_file))

        assert result.metadata == {}
        assert "Simple Document" in result.content


class TestHTMLParser:
    """Tests for the HTML document parser."""

    @pytest.mark.asyncio
    async def test_parse_html(self, sample_html_content: str, tmp_path: "Path") -> None:
        """Test basic HTML parsing with content extraction."""
        html_file = tmp_path / "test.html"
        html_file.write_text(sample_html_content)

        parser = HTMLParser()
        result = await parser.parse(str(html_file))

        assert result.file_type == "html"
        assert result.title == "Test Page"
        assert "Main Title" in result.content

    @pytest.mark.asyncio
    async def test_parse_html_clean(self, sample_html_content: str, tmp_path: "Path") -> None:
        """Test that scripts, styles, nav, and footer content is removed."""
        html_file = tmp_path / "test.html"
        html_file.write_text(sample_html_content)

        parser = HTMLParser()
        result = await parser.parse(str(html_file))

        assert "console.log" not in result.content
        assert "margin" not in result.content
        assert "Navigation bar" not in result.content
        assert "Footer content" not in result.content

    @pytest.mark.asyncio
    async def test_parse_html_sections(self, sample_html_content: str, tmp_path: "Path") -> None:
        """Test HTML section extraction based on heading tags."""
        html_file = tmp_path / "test.html"
        html_file.write_text(sample_html_content)

        parser = HTMLParser()
        result = await parser.parse(str(html_file))

        assert len(result.sections) > 0


class TestPDFParser:
    """Tests for the PDF document parser."""

    @pytest.mark.asyncio
    async def test_parse_pdf_extracts_page_text_and_metadata(self, tmp_path: "Path") -> None:
        """Test PDF parsing against a generated one-page document."""
        pytest.importorskip("fitz")
        import fitz
        from docflow.parsers.pdf import PDFParser

        pdf_file = tmp_path / "sample.pdf"
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Refund policy text from PDF")
        doc.set_metadata({"title": "Sample PDF"})
        doc.save(str(pdf_file))
        doc.close()

        parser = PDFParser()
        result = await parser.parse(str(pdf_file))

        assert result.file_type == "pdf"
        assert result.title == "Sample PDF"
        assert result.metadata["page_count"] == 1
        assert "Refund policy text" in result.content
        assert result.sections[0].title == "Page 1"


class TestDocxParser:
    """Tests for the DOCX document parser."""

    @pytest.mark.asyncio
    async def test_parse_docx_extracts_headings_paragraphs_and_tables(self, tmp_path: "Path") -> None:
        """Test DOCX parsing against a generated document."""
        pytest.importorskip("docx")
        from docx import Document
        from docflow.parsers.docx import DocxParser

        docx_file = tmp_path / "sample.docx"
        doc = Document()
        doc.add_heading("Vendor Requirements", level=1)
        doc.add_paragraph("Vendors must provide insurance certificates.")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Control"
        table.cell(0, 1).text = "Required"
        table.cell(1, 0).text = "MFA"
        table.cell(1, 1).text = "Yes"
        doc.save(str(docx_file))

        parser = DocxParser()
        result = await parser.parse(str(docx_file))

        assert result.file_type == "docx"
        assert result.title == "Vendor Requirements"
        assert result.metadata["table_count"] == 1
        assert "insurance certificates" in result.content
        assert "MFA | Yes" in result.content


class TestCSVParser:
    """Tests for the CSV document parser."""

    @pytest.mark.asyncio
    async def test_parse_csv(self, sample_csv_content: str, tmp_path: "Path") -> None:
        """Test CSV parsing with header extraction and row chunking."""
        csv_file = tmp_path / "test.csv"
        csv_file.write_text(sample_csv_content)

        parser = CSVParser()
        result = await parser.parse(str(csv_file))

        assert result.file_type == "csv"
        assert result.metadata["columns"] == ["name", "department", "salary", "startDate"]
        assert result.metadata["row_count"] == 6
        assert len(result.sections) > 0

    @pytest.mark.asyncio
    async def test_parse_csv_delimiter_detection(self, tmp_path: "Path") -> None:
        """Test automatic delimiter detection for semicolon-separated files."""
        csv_file = tmp_path / "test.csv"
        csv_file.write_text("a;b;c\n1;2;3\n4;5;6")

        parser = CSVParser()
        result = await parser.parse(str(csv_file))

        assert result.metadata["delimiter"] == ";"

    @pytest.mark.asyncio
    async def test_parse_csv_empty(self, tmp_path: "Path") -> None:
        """Test parsing an empty CSV file."""
        csv_file = tmp_path / "empty.csv"
        csv_file.write_text("")

        parser = CSVParser()
        result = await parser.parse(str(csv_file))

        assert result.content == ""
